from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "docs" / "manuales" / "Manual_usuario_Administracion_de_Alquileres_ONO_Prop.docx"
LOGO = ROOT / "public" / "assets" / "img" / "logoONOProp3.png"

NAVY = "071A3A"
BLUE = "0B4AA2"
PRIMARY = "0D6EFD"
GREEN = "16875D"
GOLD = "9A6700"
RED = "B42318"
INK = "1F2937"
MUTED = "6B7280"
LIGHT_BLUE = "EEF5FF"
LIGHT_GREEN = "EAF8F0"
LIGHT_GOLD = "FFF7E0"
LIGHT_RED = "FDECEC"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(value):
    return RGBColor.from_string(value)


def set_run_font(run, size=None, color=INK, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DCE3EC", size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths, indent=TABLE_INDENT_DXA):
    if sum(widths) != CONTENT_DXA:
        raise ValueError(f"Las columnas deben sumar {CONTENT_DXA} DXA: {widths}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cannot_split = OxmlElement("w:cantSplit")
    tr_pr.append(cannot_split)


def set_paragraph_border_bottom(paragraph, color="D7DBE2", size=8, space=6):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, NAVY, 10, 5),
    }
    for name, (size, color, before, after) in tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_numbering_definition(doc, num_format, marker, left=540, hanging=270, font="Calibri"):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(item.get(qn("w:abstractNumId"))) for item in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_ids = [int(item.get(qn("w:numId"))) for item in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), num_format)
    level.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), marker)
    level.append(text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left))
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    run_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), font)
    fonts.set(qn("w:hAnsi"), font)
    run_pr.append(fonts)
    level.append(run_pr)
    abstract.append(level)
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(numbering.index(first_num), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def clone_numbering_instance(doc, source_num_id):
    numbering = doc.part.numbering_part.element
    source = next(
        item for item in numbering.findall(qn("w:num"))
        if int(item.get(qn("w:numId"))) == int(source_num_id)
    )
    abstract_id = source.find(qn("w:abstractNumId")).get(qn("w:val"))
    num_ids = [int(item.get(qn("w:numId"))) for item in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), abstract_id)
    num.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25


def add_list_item(doc, text, num_id, bold_prefix=""):
    paragraph = doc.add_paragraph()
    apply_numbering(paragraph, num_id)
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        set_run_font(run, bold=True)
        run = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(run)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_rich_paragraph(doc, parts, after=6, keep=False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.keep_with_next = keep
    for part in parts:
        if isinstance(part, str):
            text, bold, color, italic = part, False, INK, False
        else:
            text = part.get("text", "")
            bold = part.get("bold", False)
            color = part.get("color", INK)
            italic = part.get("italic", False)
        run = paragraph.add_run(text)
        set_run_font(run, bold=bold, color=color, italic=italic)
    return paragraph


def add_callout(doc, title, body, kind="info"):
    palette = {
        "info": (LIGHT_BLUE, BLUE),
        "success": (LIGHT_GREEN, GREEN),
        "warning": (LIGHT_GOLD, GOLD),
        "danger": (LIGHT_RED, RED),
        "neutral": (LIGHT_GRAY, NAVY),
    }
    fill, accent = palette[kind]
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    set_table_borders(table, color=accent, size=8)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=130, start=180, bottom=130, end=180)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(title)
    set_run_font(run, bold=True, color=accent)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(body)
    set_run_font(run, size=10.5, color=INK)
    prevent_row_split(table.rows[0])
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_table(doc, headers, rows, widths, alignments=None):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, "E8EEF5")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run_font(run, size=9.5, bold=True, color=NAVY)
    for row_values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            if alignments:
                p.alignment = alignments[index]
            run = p.add_run(str(value))
            set_run_font(run, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_section_heading(doc, text, level=1, page_break=False):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    if page_break:
        paragraph.paragraph_format.page_break_before = True
    return paragraph


def configure_section(section):
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    left = p.add_run("ONO PROP")
    set_run_font(left, size=8.5, bold=True, color=BLUE)
    p.add_run("\t")
    right = p.add_run("MANUAL DE USUARIO · ADMINISTRACIÓN DE ALQUILERES")
    set_run_font(right, size=8, color=MUTED)
    set_paragraph_border_bottom(p, color="D7DBE2", size=6, space=4)

    footer = section.footer
    p = footer.paragraphs[0]
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    left = p.add_run("onoprop.com · Versión 1.0 · Agosto 2026")
    set_run_font(left, size=8, color=MUTED)
    p.add_run("\t")
    right = p.add_run("Página ")
    set_run_font(right, size=8, color=MUTED)
    add_field(p, "PAGE")


def add_cover(doc):
    banner = doc.add_table(rows=1, cols=1)
    set_table_geometry(banner, [CONTENT_DXA])
    set_table_borders(banner, color=NAVY, size=4)
    cell = banner.cell(0, 0)
    set_cell_shading(cell, NAVY)
    set_cell_margins(cell, top=260, start=260, bottom=260, end=260)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo = p.add_run().add_picture(str(LOGO), width=Inches(4.7))
    logo._inline.docPr.set("title", "ONO Prop")
    logo._inline.docPr.set("descr", "Logotipo de ONO Prop")

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(48)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("GUÍA OPERATIVA")
    set_run_font(run, size=10.5, bold=True, color=PRIMARY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("Administración de alquileres")
    set_run_font(run, size=30, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    run = p.add_run("Manual de usuario para inmobiliarias")
    set_run_font(run, size=15, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(80)
    run = p.add_run("Versión 1.0 · 10 de agosto de 2026")
    set_run_font(run, size=10.5, color=MUTED, italic=True)

    add_callout(
        doc,
        "Objetivo del manual",
        "Guiar al equipo de la inmobiliaria desde la carga de las partes y el contrato hasta el cobro, la liquidación al locador, la cuenta corriente y, cuando esté habilitada, la facturación electrónica.",
        "neutral",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Documento operativo. No reemplaza asesoramiento jurídico, contable ni impositivo.")
    set_run_font(run, size=9, color=MUTED, italic=True)


def build_manual():
    doc = Document()
    setup_styles(doc)
    configure_section(doc.sections[0])
    bullet_id = add_numbering_definition(doc, "bullet", "•")
    number_id = add_numbering_definition(doc, "decimal", "%1.")
    check_id = add_numbering_definition(doc, "bullet", "☐", font="Segoe UI Symbol")

    add_cover(doc)
    doc.add_page_break()

    add_section_heading(doc, "Cómo usar este manual", 1)
    add_rich_paragraph(doc, [
        "Las instrucciones siguen las etiquetas visibles en ONO Prop. Cuando se menciona un botón o estado, aparece en ",
        {"text": "negrita", "bold": True},
        ". La disponibilidad de una acción depende del módulo contratado y del rol asignado al usuario.",
    ])
    add_callout(
        doc,
        "Antes de comenzar",
        "Verificá que la inmobiliaria correcta esté seleccionada en la barra de navegación. Los contratos, personas y movimientos pertenecen siempre a la inmobiliaria activa.",
        "warning",
    )

    add_section_heading(doc, "Ruta rápida recomendada", 2)
    quick_steps = [
        "Cargar el inmueble, aunque todavía no esté publicado.",
        "Registrar locadores, locatarios y garantes en Personas.",
        "Crear el contrato como Borrador, revisar todos los datos y luego activarlo.",
        "Controlar las obligaciones generadas y registrar cada cobro en el período correcto.",
        "Emitir el recibo del locatario, cargar gastos y preparar la liquidación.",
        "Registrar el pago al locador y confirmar posteriormente la recepción.",
        "Revisar la cuenta corriente y los fondos pendientes de entrega.",
        "Si corresponde, preparar la documentación fiscal y emitir sólo después de revisar la vista previa.",
    ]
    quick_number_id = clone_numbering_instance(doc, number_id)
    for item in quick_steps:
        add_list_item(doc, item, quick_number_id)

    add_section_heading(doc, "Contenido", 2)
    contents = [
        "1. Acceso, permisos y panel principal",
        "2. Personas y partes",
        "3. Creación y edición del contrato",
        "4. Obligaciones, ajustes y gastos",
        "5. Cobros, recibos y cancelaciones externas",
        "6. Liquidaciones y pagos al locador",
        "7. Cuentas corrientes y fondos pendientes",
        "8. Facturación ARCA: HOMO y Producción",
        "9. Rectificaciones, archivo y controles periódicos",
        "10. Estados, preguntas frecuentes y soporte",
    ]
    for item in contents:
        add_list_item(doc, item, bullet_id)

    add_section_heading(doc, "1. Acceso, permisos y panel principal", 1, page_break=True)
    add_section_heading(doc, "1.1 Ingreso al módulo", 2)
    add_rich_paragraph(doc, [
        "Iniciá sesión, abrí el menú administrativo y elegí ",
        {"text": "Administración de alquileres", "bold": True},
        ". Si administrás más de una inmobiliaria, confirmá la inmobiliaria activa antes de cargar o modificar datos.",
    ])
    add_callout(
        doc,
        "Si el módulo no aparece",
        "La inmobiliaria puede no tener el servicio habilitado o tu usuario puede no tener permiso. Consultá al administrador interno o a ONO Prop.",
        "info",
    )

    add_section_heading(doc, "1.2 Permisos internos", 2)
    add_table(
        doc,
        ["Rol", "Puede consultar", "Puede gestionar"],
        [
            ["Administrador", "Contratos, movimientos, recibos y cuentas", "Personas, contratos, cobros, gastos y liquidaciones"],
            ["Editor", "Contratos, movimientos, recibos y cuentas", "Personas, contratos, cobros, gastos y liquidaciones"],
            ["Lector", "Contratos, movimientos, recibos y cuentas", "No puede registrar ni rectificar movimientos"],
        ],
        [1800, 3240, 4320],
    )
    add_rich_paragraph(doc, [
        "La emisión fiscal real requiere, además, una habilitación expresa de ONO Prop. Ver el capítulo 8.",
    ])

    add_section_heading(doc, "1.3 Panel principal", 2)
    add_rich_paragraph(doc, [
        "El panel resume ", {"text": "Contratos activos", "bold": True}, ", ",
        {"text": "Obligaciones vencidas", "bold": True}, ", ",
        {"text": "Cobrado este período", "bold": True}, " y ",
        {"text": "Ajustes en 45 días", "bold": True}, ".",
    ])
    people_number_id = clone_numbering_instance(doc, number_id)
    for item in [
        "Usá Buscar contrato para localizar por inmueble, dirección, locador o locatario.",
        "Filtrá por Borrador, Activo, Finalizado o Rescindido.",
        "Elegí Gestionar para operar el contrato o Editar contrato para cambiar sus condiciones.",
        "Accedé a Personas y Cuentas de locadores desde los botones superiores.",
    ]:
        add_list_item(doc, item, bullet_id)

    add_section_heading(doc, "2. Personas y partes", 1)
    add_section_heading(doc, "2.1 Registrar una persona", 2)
    for item in [
        "Entrá en Personas y completá Nombre o razón social.",
        "Elegí Persona humana o Persona jurídica.",
        "Ingresá DNI o CUIT y la condición frente al IVA. Estos datos son especialmente importantes para ARCA.",
        "Marcá al menos un rol: Locador, Locatario o Garante. Una misma persona puede tener más de un rol.",
        "Completá email, teléfono, domicilio y, para locadores, CBU, CVU o alias.",
        "Guardá. La persona aparecerá en el directorio y en los selectores de contratos.",
    ]:
        add_list_item(doc, item, people_number_id)
    add_callout(
        doc,
        "Dato histórico",
        "Al vincular una persona, el contrato conserva una copia de su identificación. Si después editás el directorio, revisá el contrato cuando necesites reflejar el cambio en ese vínculo histórico.",
        "info",
    )

    add_section_heading(doc, "2.2 Editar o archivar", 2)
    add_rich_paragraph(doc, [
        "Usá ", {"text": "Editar", "bold": True}, " para corregir datos. Usá ",
        {"text": "Archivar", "bold": True}, " cuando la persona ya no deba aparecer en nuevas cargas. Archivar no elimina su identificación de los contratos existentes.",
    ])

    add_section_heading(doc, "3. Creación y edición del contrato", 1)
    add_section_heading(doc, "3.1 Requisitos previos", 2)
    adjustment_number_id = clone_numbering_instance(doc, number_id)
    for item in [
        "El inmueble debe estar cargado en la inmobiliaria, publicado o no.",
        "Debe existir al menos un locador y un locatario activos en Personas.",
        "El usuario debe tener permiso para gestionar alquileres.",
    ]:
        add_list_item(doc, item, check_id)

    add_section_heading(doc, "3.2 Completar el formulario", 2)
    contract_steps = [
        ("Inmueble y estado", "Seleccioná el inmueble administrado. Para revisar antes de operar, comenzá con estado Borrador."),
        ("Partes", "Elegí uno o más locadores y locatarios. Los garantes son opcionales. El sistema impide que la misma persona sea locador y locatario en un mismo contrato."),
        ("Vigencia y valores", "Ingresá inicio, finalización, fecha de firma, día mensual de vencimiento, moneda, alquiler inicial y depósito."),
        ("Honorarios", "Podés establecer un porcentaje mensual, un importe fijo o ambos. El cálculo nunca supera el monto cobrado."),
        ("Actualización", "Elegí actualización manual, porcentaje fijo, índice de referencia o fórmula contractual. Para índice o fórmula, el nuevo importe se confirma manualmente."),
        ("Documentación y notas", "Registrá conceptos incluidos, mora pactada, enlace al contrato digitalizado y notas internas."),
    ]
    for title, body in contract_steps:
        p = doc.add_paragraph(style="Heading 3")
        p.add_run(title)
        add_rich_paragraph(doc, [body])

    add_callout(
        doc,
        "Cuándo activar",
        "Activá el contrato después de verificar inmueble, partes, vigencia, alquiler, vencimiento y actualización. Al activarlo se generan obligaciones mensuales; luego la automatización mantiene cubiertos los próximos 45 días.",
        "success",
    )

    add_section_heading(doc, "3.3 Estados del contrato", 2)
    add_table(
        doc,
        ["Estado", "Uso recomendado"],
        [
            ["Borrador", "Carga y revisión antes de comenzar la administración."],
            ["Activo", "Contrato vigente con generación y gestión de obligaciones."],
            ["Finalizado", "La vigencia terminó normalmente."],
            ["Rescindido", "El contrato terminó anticipadamente."],
        ],
        [1800, 7560],
    )

    add_section_heading(doc, "3.4 Editar un contrato vigente", 2)
    add_rich_paragraph(doc, [
        "Desde el panel elegí ", {"text": "Editar contrato", "bold": True}, ". Después de guardar, revisá los períodos. Si la vigencia cambió, el bloque ARCA puede mostrar ",
        {"text": "Sincronizar períodos ahora", "bold": True}, ".",
    ])
    add_callout(
        doc,
        "Protección del historial",
        "Los períodos sin actividad pueden sincronizarse o apartarse. Los que ya tienen pagos, cancelaciones externas o gestión fiscal quedan preservados y requieren revisión manual.",
        "warning",
    )

    add_section_heading(doc, "4. Obligaciones, ajustes y gastos", 1)
    add_section_heading(doc, "4.1 Generar obligaciones", 2)
    add_rich_paragraph(doc, [
        "En el detalle del contrato, el bloque ", {"text": "Generar obligaciones", "bold": True}, " permite elegir una fecha y presionar ",
        {"text": "Generar hasta fecha", "bold": True}, ". Los períodos existentes no se duplican.",
    ])
    add_rich_paragraph(doc, [
        "Los contratos activos también se procesan automáticamente todos los días. El sistema procura dejar creadas las obligaciones comprendidas dentro de los próximos 45 días.",
    ])

    add_section_heading(doc, "4.2 Leer una obligación", 2)
    add_table(
        doc,
        ["Dato", "Interpretación"],
        [
            ["Total", "Alquiler del período más otros cargos registrados."],
            ["Cobrado por la inmobiliaria", "Suma de pagos activos imputados al período."],
            ["Cerrado fuera de gestión", "Importe cerrado sin simular una cobranza de la inmobiliaria."],
            ["Saldo operativo", "Importe todavía pendiente dentro de la gestión."],
            ["Neto locador", "Cobrado menos honorarios y gastos a cargo del locador."],
        ],
        [2500, 6860],
    )

    add_section_heading(doc, "4.3 Confirmar un nuevo alquiler", 2)
    for item in [
        "Ingresá la fecha desde la cual rige el ajuste.",
        "Ingresá el nuevo importe.",
        "Documentá el índice, la fuente o el cálculo en Notas del ajuste.",
        "Presioná Registrar y revisá las obligaciones afectadas.",
    ]:
        add_list_item(doc, item, adjustment_number_id)
    add_callout(
        doc,
        "Índice o fórmula",
        "El sistema registra la regla y genera alertas, pero exige confirmar el importe efectivamente aplicado para conservar trazabilidad.",
        "info",
    )

    add_section_heading(doc, "4.4 Otros cargos y gastos", 2)
    add_rich_paragraph(doc, [
        "Usá ", {"text": "Otros cargos del período", "bold": True}, " para incorporar conceptos que aumentan la obligación. Usá ",
        {"text": "Registrar gasto", "bold": True}, " para registrar egresos y asignarlos al locador, locatario o inmobiliaria.",
    ])
    add_rich_paragraph(doc, [
        "Sólo los gastos asignados ", {"text": "A cargo del locador", "bold": True}, " se descuentan del neto de su liquidación. Los demás quedan registrados sin disminuir ese neto.",
    ])

    add_section_heading(doc, "5. Cobros, recibos y cancelaciones externas", 1, page_break=True)
    add_section_heading(doc, "5.1 Registrar un cobro", 2)
    payment_number_id = clone_numbering_instance(doc, number_id)
    for item in [
        "Ubicá el período correcto y presioná Registrar pago.",
        "Confirmá importe y fecha. El importe propuesto es el saldo, pero puede registrarse un pago parcial.",
        "Elegí transferencia, efectivo, depósito, tarjeta/plataforma u otro.",
        "Agregá referencia y notas cuando exista un comprobante bancario o una aclaración.",
        "Presioná Confirmar pago.",
    ]:
        add_list_item(doc, item, payment_number_id)
    add_callout(
        doc,
        "Imputación correcta",
        "Un cobro modifica saldos, recibos, liquidación y cuenta corriente. Verificá siempre contrato, período, fecha e importe antes de confirmarlo.",
        "warning",
    )

    add_section_heading(doc, "5.2 Recibo del locatario", 2)
    add_rich_paragraph(doc, [
        "Después del cobro, elegí ", {"text": "Ver recibo", "bold": True}, " y luego ",
        {"text": "Imprimir triplicado / guardar PDF", "bold": True}, ". El sistema genera:",
    ])
    owner_payment_number_id = clone_numbering_instance(doc, number_id)
    for item in [
        "Original para el locatario.",
        "Duplicado para el locador.",
        "Triplicado para la inmobiliaria.",
    ]:
        add_list_item(doc, item, bullet_id)
    add_rich_paragraph(doc, [
        "El recibo expresa la moneda y el monto en letras mayúsculas, seguido por la cifra entre paréntesis.",
    ])

    add_section_heading(doc, "5.3 Cancelación externa", 2)
    add_rich_paragraph(doc, [
        "Usá ", {"text": "Cancelación externa", "bold": True}, " cuando el período no fue cobrado por la inmobiliaria: por ejemplo, era anterior al inicio de la administración, el locador cobró directamente o intervino un tercero.",
    ])
    add_rich_paragraph(doc, [
        "Esta acción lleva el saldo operativo a cero, pero ",
        {"text": "no genera cobranza, recibo, honorario ni liquidación", "bold": True, "color": RED},
        ". Puede revertirse mediante ", {"text": "Reabrir período", "bold": True}, ".",
    ])

    add_section_heading(doc, "5.4 Facturación externa", 2)
    add_rich_paragraph(doc, [
        "La opción ", {"text": "Factura externa", "bold": True}, " indica que la gestión fiscal se realizó fuera de ONO Prop y evita una nueva preparación ARCA para ese período.",
    ])
    for item in [
        "Si conocés los datos, elegí tipo, punto de venta, número, fecha, importe y opcionalmente CAE.",
        "Si no disponés del comprobante, elegí Sin datos del comprobante.",
        "Usá Editar datos o Quitar marca sólo si corresponde corregir la clasificación.",
    ]:
        add_list_item(doc, item, bullet_id)
    add_callout(
        doc,
        "No confundir",
        "Cancelación externa describe quién administró el cobro. Facturación externa describe dónde se gestionó el comprobante fiscal. Son registros distintos.",
        "info",
    )

    add_section_heading(doc, "6. Liquidaciones y pagos al locador", 1)
    add_section_heading(doc, "6.1 Preparar la liquidación", 2)
    add_rich_paragraph(doc, [
        "Cuando exista al menos un cobro, presioná ", {"text": "Liquidar", "bold": True}, ". El cálculo toma el cobrado por la inmobiliaria y deduce:",
    ])
    for item in [
        "Honorarios porcentuales y/o fijos del período.",
        "Gastos registrados a cargo del locador.",
    ]:
        add_list_item(doc, item, bullet_id)
    add_rich_paragraph(doc, [
        "Si luego cambia un cobro o gasto y la liquidación sigue en borrador, usá ", {"text": "Recalcular", "bold": True}, ".",
    ])

    add_section_heading(doc, "6.2 Registrar el pago al locador", 2)
    for item in [
        "Presioná Registrar pago al locador.",
        "Ingresá fecha, medio de pago, referencia y observaciones.",
        "Confirmá para generar el recibo de liquidación.",
        "Abrí Recibo locador e imprimí el duplicado: original para el locador y duplicado para la inmobiliaria.",
    ]:
        add_list_item(doc, item, owner_payment_number_id)
    add_callout(
        doc,
        "Estado intermedio",
        "Registrar el pago indica que el dinero salió de la inmobiliaria. La liquidación queda como Pago registrado · recepción pendiente hasta que exista confirmación del locador.",
        "warning",
    )

    add_section_heading(doc, "6.3 Confirmar la recepción", 2)
    add_rich_paragraph(doc, [
        "Presioná ", {"text": "Confirmar recepción", "bold": True}, " y registrá fecha, medio de confirmación, referencia y observaciones. Los medios disponibles incluyen recibo firmado, confirmación escrita, acreditación bancaria, confirmación verbal u otro medio documentado.",
    ])
    add_rich_paragraph(doc, [
        "Al confirmar, el estado cambia a ", {"text": "Recepción confirmada por el locador", "bold": True, "color": GREEN}, ".",
    ])

    add_section_heading(doc, "7. Cuentas corrientes y fondos pendientes", 1)
    add_section_heading(doc, "7.1 Cuenta de un contrato", 2)
    add_rich_paragraph(doc, [
        "Desde el contrato elegí ", {"text": "Cuenta corriente", "bold": True}, ". Los créditos representan cobros administrados a favor del locador. Los débitos incluyen honorarios, gastos a su cargo y pagos realizados.",
    ])
    add_rich_paragraph(doc, [
        "Un pago con recepción pendiente reduce el saldo porque el dinero ya salió de la inmobiliaria, pero conserva la advertencia hasta la confirmación.",
    ])

    add_section_heading(doc, "7.2 Vista consolidada", 2)
    add_rich_paragraph(doc, [
        "En el panel principal presioná ", {"text": "Cuentas de locadores", "bold": True}, ". Allí se muestran fondos pendientes de entrega, cantidad de locadores con saldo y cuenta consolidada por locador.",
    ])
    fiscal_draft_number_id = clone_numbering_instance(doc, number_id)
    for item in [
        "Elegí Ver detalle o seleccioná un locador.",
        "Si existen contratos en distintas monedas, elegí la moneda. Los saldos no se mezclan.",
        "Usá Imprimir cuenta seleccionada para entregar o archivar el estado de cuenta.",
    ]:
        add_list_item(doc, item, bullet_id)

    add_section_heading(doc, "8. Facturación ARCA: HOMO y Producción", 1, page_break=True)
    add_callout(
        doc,
        "Alcance actual",
        "Esta etapa admite Factura C de servicios para contratos en pesos argentinos (ARS). La emisión real sólo aparece en cuentas expresamente habilitadas y con un perfil fiscal activo.",
        "warning",
    )
    add_section_heading(doc, "8.1 Diferencia entre ambientes", 2)
    add_table(
        doc,
        ["Ambiente", "Efecto", "Uso"],
        [
            ["HOMO", "Genera un CAE de homologación sin valor fiscal real", "Pruebas de datos, fechas, punto de venta y conexión"],
            ["PROD", "Genera una factura real con CAE y numeración fiscal", "Sólo después de una revisión definitiva"],
        ],
        [1500, 3660, 4200],
    )

    add_section_heading(doc, "8.2 Preparar el borrador fiscal", 2)
    for item in [
        "En Comprobantes ARCA, elegí Preparar junto al período.",
        "Revisá perfil emisor y locatario receptor.",
        "Confirmá nombre o razón social, documento y condición frente al IVA.",
        "Revisá fecha del comprobante, importe en ARS, período de servicio y vencimiento.",
        "Si modificás importe, fechas o vencimiento respecto del contrato, documentá el Motivo de la excepción.",
        "Presioná Validar y guardar borrador.",
    ]:
        add_list_item(doc, item, fiscal_draft_number_id)
    add_callout(
        doc,
        "Fechas de servicios",
        "ARCA admite una fecha de comprobante comprendida entre 10 días antes y 10 días después de la solicitud. El vencimiento para el pago no puede ser anterior a la fecha del comprobante.",
        "info",
    )

    add_section_heading(doc, "8.3 Probar en homologación", 2)
    add_rich_paragraph(doc, [
        "Presioná ", {"text": "Solicitar CAE HOMO", "bold": True}, ". El sistema advierte expresamente que no es una factura real. Si ARCA autoriza, el período mostrará ",
        {"text": "Autorizado HOMO", "bold": True}, " y podrás abrir ", {"text": "Ver HOMO", "bold": True}, ".",
    ])

    add_section_heading(doc, "8.4 Preparar y emitir en Producción", 2)
    production_number_id = clone_numbering_instance(doc, number_id)
    for item in [
        "Desde un borrador validado, presioná Preparar PROD o Actualizar PROD.",
        "La preparación consulta la numeración real, pero todavía no solicita CAE ni reserva número.",
        "Revisá emisor, CUIT, receptor, período, importe, fechas y número estimado en Vistas previas de Producción.",
        "Presioná Emitir factura real sólo cuando todo sea definitivo.",
        "Confirmá la advertencia y escribí EMITIR en la segunda confirmación.",
        "Esperá el resultado. Si fue autorizada, abrí Ver factura PROD y guardá la representación con QR.",
    ]:
        add_list_item(doc, item, production_number_id)
    add_callout(
        doc,
        "Acción irreversible",
        "Una factura real autorizada no se elimina ni se reclasifica como externa. Si contiene un error, no alteres el período: consultá al responsable fiscal y a ONO Prop para definir el comprobante de ajuste correspondiente.",
        "danger",
    )
    add_callout(
        doc,
        "Respuesta a reconciliar",
        "Si aparece este estado, no prepares otra factura. Presioná Reconciliar emisión: el sistema consulta si ARCA ya autorizó el número propuesto y evita duplicados.",
        "warning",
    )

    add_section_heading(doc, "8.5 Facturación de períodos renegociados", 2)
    add_rich_paragraph(doc, [
        "El borrador permite editar importe, servicio y vencimiento sin modificar la obligación contractual. Usá esta excepción sólo para reflejar lo efectivamente acordado y documentá el motivo, por ejemplo: importe y fecha renegociados con el locatario.",
    ])

    add_section_heading(doc, "9. Rectificaciones, archivo y controles", 1, page_break=True)
    add_section_heading(doc, "9.1 Anular un cobro", 2)
    add_rich_paragraph(doc, [
        "Junto al pago presioná ", {"text": "Anular cobro", "bold": True}, " e ingresá un motivo. El movimiento no se borra: queda marcado como anulado, se recalculan saldo y estado, y la cuenta corriente deja de computarlo.",
    ])
    add_callout(
        doc,
        "Liquidación relacionada",
        "No puede anularse un cobro si la liquidación ya figura pagada o recibida. Primero debe rectificarse el circuito del locador en orden inverso.",
        "warning",
    )

    add_section_heading(doc, "9.2 Rectificar el circuito del locador", 2)
    correction_number_id = clone_numbering_instance(doc, number_id)
    for item in [
        "Si la recepción fue confirmada por error, usá Rectificar recepción e indicá el motivo.",
        "La liquidación volverá a Pago registrado · recepción pendiente.",
        "Si también fue erróneo el pago, usá Anular pago al locador e indicá el motivo.",
        "La liquidación volverá a borrador y conservará el historial de correcciones.",
    ]:
        add_list_item(doc, item, correction_number_id)

    add_section_heading(doc, "9.3 Archivar el contrato", 2)
    add_rich_paragraph(doc, [
        "Usá ", {"text": "Archivar", "bold": True}, " cuando el contrato ya no deba aparecer en la gestión habitual. El archivo no elimina obligaciones, pagos, recibos ni liquidaciones.",
    ])

    add_section_heading(doc, "9.4 Rutina operativa recomendada", 2)
    add_table(
        doc,
        ["Frecuencia", "Control"],
        [
            ["Diaria", "Revisar vencidas, registrar cobros, emitir recibos y comprobar mensajes de error."],
            ["Semanal", "Revisar fondos pendientes, liquidaciones pagadas sin recepción y ajustes próximos."],
            ["Mensual", "Conciliar cuentas de locadores, controlar gastos, períodos y contratos que terminan."],
            ["Antes de PROD", "Revisar perfil, receptor, importe, fechas, numeración estimada y respaldo documental."],
        ],
        [1500, 7860],
    )

    add_section_heading(doc, "10. Estados, preguntas frecuentes y soporte", 1, page_break=True)
    add_section_heading(doc, "10.1 Estados de las obligaciones", 2)
    add_table(
        doc,
        ["Estado", "Significado"],
        [
            ["Pendiente", "Todavía no tiene cobros y no está vencida."],
            ["Pago parcial", "Tiene cobros, pero conserva saldo antes del vencimiento."],
            ["Vencida", "Conserva saldo después del vencimiento."],
            ["Pagada", "El saldo operativo llegó a cero mediante cobros de la inmobiliaria."],
            ["Cancelación externa", "El saldo fue cerrado fuera de la gestión de la inmobiliaria."],
        ],
        [2100, 7260],
    )

    add_section_heading(doc, "10.2 Estados de las liquidaciones", 2)
    add_table(
        doc,
        ["Estado", "Acción siguiente"],
        [
            ["Liquidación preparada", "Revisar cálculo y registrar pago al locador."],
            ["Requiere recálculo", "Recalcular porque cambió un movimiento relacionado."],
            ["Pago registrado · recepción pendiente", "Obtener respaldo y confirmar recepción."],
            ["Recepción confirmada", "Circuito administrativo finalizado."],
        ],
        [3300, 6060],
    )

    add_section_heading(doc, "10.3 Preguntas frecuentes", 2)
    faqs = [
        ("No aparece un inmueble al crear el contrato.", "Debe estar cargado dentro de la inmobiliaria activa. Puede permanecer sin publicar."),
        ("Una persona no aparece como locador o locatario.", "Revisá que esté activa y tenga marcado el rol correspondiente."),
        ("El contrato editado conserva períodos anteriores.", "Usá Sincronizar períodos ahora. Los períodos con actividad se preservan para revisión."),
        ("No puedo liquidar.", "Debe existir al menos un cobro activo y el período no debe estar cerrado externamente. Una liquidación pagada o recibida no puede recalcularse."),
        ("No aparece un perfil ARCA.", "La inmobiliaria no tiene un perfil fiscal activo o tu cuenta no está habilitada. Contactá a ONO Prop."),
        ("ARCA rechaza la fecha.", "Verificá el rango de 10 días respecto de la solicitud y que el vencimiento no sea anterior al comprobante."),
        ("El contrato está en USD o EUR.", "La facturación ARCA integrada actual admite únicamente contratos en ARS."),
        ("La impresión no muestra el recibo.", "Abrí el recibo desde Ver recibo o Recibo locador y usá el botón de impresión de esa pantalla. Revisá la vista previa antes de guardar PDF."),
    ]
    for question, answer in faqs:
        p = doc.add_paragraph(style="Heading 3")
        p.add_run(question)
        add_rich_paragraph(doc, [answer])

    add_section_heading(doc, "10.4 Buenas prácticas de seguridad", 2)
    for item in [
        "No compartas usuarios ni contraseñas entre operadores.",
        "Asigná el rol Lector a quienes sólo deban consultar.",
        "No incluyas contraseñas ni secretos fiscales en notas del contrato.",
        "Guardá referencias de transferencias, confirmaciones y recibos firmados.",
        "Ante una duda fiscal productiva, detené la emisión y conservá el mensaje exacto.",
    ]:
        add_list_item(doc, item, check_id)

    add_callout(
        doc,
        "Soporte ONO Prop",
        "Al solicitar ayuda, informá la inmobiliaria, el contrato, el período, la acción realizada y el mensaje completo. No envíes certificados, claves privadas ni contraseñas. Contacto: contacto@onoprop.com.",
        "neutral",
    )

    add_section_heading(doc, "Control de versión", 2)
    add_table(
        doc,
        ["Versión", "Fecha", "Alcance"],
        [["1.0", "10/08/2026", "Primera edición: circuito completo de administración de alquileres y facturación ARCA."]],
        [1400, 1900, 6060],
    )
    add_rich_paragraph(doc, [
        {"text": "Importante: ", "bold": True},
        "la interfaz puede incorporar mejoras posteriores. Ante diferencias, prevalecen las validaciones y advertencias visibles en ONO Prop.",
    ])

    core = doc.core_properties
    core.title = "Manual de usuario - Administración de alquileres"
    core.subject = "Guía operativa del módulo de Administración de alquileres de ONO Prop"
    core.author = "ONO Prop"
    core.keywords = "ONO Prop, alquileres, inmobiliaria, contratos, cobros, ARCA"
    core.comments = "Versión 1.0 - Agosto 2026"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_manual()
